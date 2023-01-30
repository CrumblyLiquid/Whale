# Umožnění anotování vracení sama sebe (resp. typu třídy): class Monke: def fun() -> Monke: ...
from __future__ import annotations
from io import TextIOWrapper
import json
from os import path as ospath
import sys # Argumenty

# Importy pro konkrétní soubory
import pandas as pd
# Potřebuje pyexcel, pyxecel-xlsx, pyexcel-ods, ...
# Více na: https://github.com/pyexcel/pyexcel#why-did-not-i-see-above-benefit
# a https://github.com/pyexcel/pyexcel#available-plugins
import pyexcel
# Dokumenty
from docx import Document
from odf import text as odftext, teletype
from odf.style import Style, TextProperties
from odf import opendocument

# Vrátí jméno souboru z cesty
def name(filename: str) -> str:
    return ospath.basename(filename)

# Vrátí typ souboru z jeho jména/cesty
def get_extension(filename: str) -> str:
    filename = name(filename)
    return filename.split(".")[-1]

class UnsupportedFormatException(Exception):
    pass

# Přechodná struktura pro převádění mezi typy souborů
class Package:
    name: str
    inputs: dict
    words: list[list[str]]

    def __init__(self, name: str, inputs: dict, words: list[list[str]]) -> None:
        self.name = name
        self.inputs = inputs
        self.words = words

    # Vrátí příslušnou funkci podle rozšíření
    def fn_from_ext(ext: str) -> str:
        supported = {
            "table": ["xlsx", "xls", "ods"],
            "document": ["docx", "doc", "odt"],
            "text": ["txt"],
            "json": ["json"],
        }

        for fn in supported.keys():
            if ext in supported[fn]:
                return fn
        return "unsupported"

    def size(self) -> int:
        return len(self.words)

    # Vytvoří Package z DataFrame (dá se použít pro xlsx, ods, csv, json, ...)
    def from_dataframe(df: pd.DataFrame) -> Package:
        words = []
        inputs = df.columns.to_list()
        for (label, series) in df.iterrows():
            if(series.isnull().values.any() == False):
                words.append(list(series.values))

        return Package(name="dataframe", inputs=inputs, words=words)

    # from_x funkce přijmou file handle a extrahují data do třídy Package
    # to_x funkce přijmou file handle a uloží do něj data v příslušném formátu
    # samotné získání file handlu je přenecháno vlastní interpretaci

    def from_text(filename: str) -> Package:
        words = []
        inputs = []
        with open(filename, "r", encoding="utf-8") as file:
            for line in file.readlines():
                # Musíme řádek zbavit \n -> strip()
                if("|" in line):
                    inputs = line.strip().split(" | ")
                elif("-" in line):
                    # Multiword = více slov v jiných jazycích/tvarech
                    multiword = line.split("-")
                    stripped_multiword = []
                    for word in multiword:
                        stripped_multiword.append(word.strip())
                    words.append(stripped_multiword)

        return Package(ospath.basename(filename), inputs=inputs, words=words)

    def to_text(self, filename: str) -> None:
        with open(file=filename, mode="w", encoding="utf-8") as file:
            # list pospojuji s \n -> nemusím na to pamatovat jako u stringu
            result = []

            result.append(" | ".join(self.inputs))

            # Vložit příslušně spojená slova
            for multiword in self.words:
                word_str = " - ".join(multiword)
                result.append(word_str)

            # Výsledný string na zapsání
            res = "\n".join(result)

            # Zapsat do souboru
            file.write(res)

    def from_table(filename: str) -> Package:
        # xlsx, ods - potřeba openpyxl nebo jiný engine!
        data = pd.read_excel(filename, sheet_name=0)
        return Package.from_dataframe(data)

    # Neřeší šířku tabulky!
    def to_table(self, filename: str) -> None:
        data = [self.inputs] + self.words
        pyexcel.save_as(array=data, dest_file_name=filename)

    def from_document(filename: str) -> Package:
        ext = get_extension(filename)
        text = []
        if(ext == "docx"):
            doc = Document(filename)
            for paragraph in doc.paragraphs:
                text = text + paragraph.text.splitlines()
        elif(ext == "odt"):
            doc = opendocument.load(filename)
            paragraphs = doc.getElementsByType(odftext.P)
            for paragraph in paragraphs:
                text = text + teletype.extractText(paragraph).splitlines()
        else:
            raise UnsupportedFormatException(f".{ext} files are not supported yet")

        words = []
        inputs = []
        for line in text:
            # Musíme řádek zbavit \n -> strip()
            if("|" in line):
                inputs = line.strip().split(" | ")
            elif("-" in line):
                # Multiword = více slov v jiných jazycích/tvarech
                multiword = line.split("-")
                stripped_multiword = []
                for word in multiword:
                    stripped_multiword.append(word.strip())
                words.append(stripped_multiword)

        return Package(ospath.basename(filename), inputs=inputs, words=words)

    def to_document(self, filename: str) -> None:
        ext = get_extension(filename)
        inputs = " | ".join(self.inputs)
        result = []
        for multiword in self.words:
            word_str = " - ".join(multiword)
            result.append(word_str)
        words = "\n".join(result)
        if(ext == "docx"):
            doc = Document()
            paragraph = doc.add_paragraph("")
            paragraph.add_run(inputs).bold = True
            # Vložit příslušně spojená slova
            doc.add_paragraph(words)
            doc.save(filename)
        elif(ext == "odt"):
            doc = opendocument.OpenDocumentText()
            boldstyle = Style(name="BoldStyle", family="paragraph")
            boldstyle.addElement(TextProperties(attributes={"fontweight": "bold"}))
            doc.styles.addElement(boldstyle)
            inputs_paragraph = odftext.P(text=inputs, stylename=boldstyle)
            doc.text.addElement(inputs_paragraph)
            for word in result:
                words_paragraph = odftext.P(text=word)
                doc.text.addElement(words_paragraph)
            doc.save(filename)
        else:
            raise UnsupportedFormatException(f".{ext} files are not supported yet")


    # TODO: JSON Implementace podle Package structu ve Whale
    def from_json(filename: str) -> Package:
        pass

    def to_json(self, filename: str) -> None:
        # print(self.inputs)
        inputs = []
        for input in self.inputs:
            inputs.append({"name": input})

        structure = {
            "name": self.name,
            "summary": "",
            "author": "CrumblyLiquid",
            "native": self.inputs[0],
            "foreign": self.inputs[1],
            "inputs": inputs,
            "words": self.words
        }

        with open(filename, "w", encoding='utf8') as file:
            json.dump(structure, file, indent=4, ensure_ascii=False)

# Samotná konverze z jednoho formátu do druhého
if __name__ == "__main__":
    args = sys.argv[1:]
    # file = open(filename, "r", encoding="utf-8")
    from_file = args[0]
    to_file = args[1]

    # Create Package object by loading one of the supported filetypes
    pkg = Package
    from_type = Package.fn_from_ext(get_extension(from_file))
    match(from_type):
        case "text":
            pkg = Package.from_text(from_file)
        case "table":
            pkg = Package.from_table(from_file)
        case "document":
            pkg = Package.from_document(from_file)
        case "json":
            pkg = Package.from_json(from_file)

    # Save Package object by serializing to the desired filetype
    to_type = Package.fn_from_ext(get_extension(to_file))
    match(to_type):
        case "text":
            pkg.to_text(to_file)
        case "table":
            pkg.to_table(to_file)
        case "document":
            pkg.to_document(to_file)
        case "json":
            pkg.to_json(to_file)